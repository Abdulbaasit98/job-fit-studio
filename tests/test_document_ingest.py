import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from job_fit_studio.document_ingest import (extract_text, ingest_directory, ingest_document,
                                               split_into_paragraphs)


def test_split_into_paragraphs_splits_on_blank_lines():
    text = ("First paragraph about a project.\n\n"
            "Second paragraph about a skill.\n\n"
            "Third paragraph also has enough words to count.")
    paragraphs = split_into_paragraphs(text)
    assert len(paragraphs) == 3
    assert paragraphs[0] == "First paragraph about a project."


def test_split_into_paragraphs_filters_short_fragments():
    text = "SKILLS\n\nThis is a real paragraph with enough words to count as content."
    paragraphs = split_into_paragraphs(text, min_words=5)
    assert "SKILLS" not in paragraphs
    assert len(paragraphs) == 1


def test_split_into_paragraphs_collapses_internal_whitespace():
    text = "This   paragraph\nhas  weird\n\n   spacing   that should   collapse cleanly."
    paragraphs = split_into_paragraphs(text)
    assert "  " not in paragraphs[0]  # no double-spaces survive


def test_extract_text_from_txt(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("Some resume content here about my real project experience.")
    result = extract_text(f)
    assert "resume content" in result


def test_extract_text_from_md(tmp_path):
    f = tmp_path / "portfolio.md"
    f.write_text("# My Portfolio\n\nI built several real projects with real tests.")
    result = extract_text(f)
    assert "real projects" in result


def test_extract_text_from_docx(tmp_path):
    import docx
    f = tmp_path / "resume.docx"
    d = docx.Document()
    d.add_paragraph("I have real experience with Python and machine learning systems.")
    d.add_paragraph("I also built a computer vision pipeline with real automated tests.")
    d.save(str(f))

    result = extract_text(f)
    assert "Python and machine learning" in result
    assert "computer vision pipeline" in result


def test_extract_text_from_pdf(tmp_path):
    from pypdf import PdfWriter
    import io

    # Build a minimal real PDF with actual extractable text using reportlab
    # if available, otherwise skip gracefully -- pypdf alone can't easily
    # WRITE text into a new PDF (it's a reader/manipulator, not a text-PDF
    # generator), so we check for reportlab as the actual PDF-writing tool.
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        import pytest
        pytest.skip("reportlab not installed -- can't generate a real text PDF for this test")

    f = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(f))
    c.drawString(100, 750, "I have real experience building deployed ML systems.")
    c.save()

    result = extract_text(f)
    assert "deployed ML systems" in result


def test_extract_text_unsupported_type_returns_empty_string(tmp_path):
    f = tmp_path / "image.png"
    f.write_bytes(b"not a real image, just bytes")
    result = extract_text(f)
    assert result == ""


def test_ingest_document_produces_tagged_chunks(tmp_path):
    f = tmp_path / "my_resume.txt"
    f.write_text("First real paragraph about a shipped project.\n\nSecond real paragraph about a skill area.")

    chunks = ingest_document(f, kind="resume")

    assert len(chunks) == 2
    assert all(c.kind == "resume" for c in chunks)
    assert all("my_resume.txt" in c.evidence for c in chunks)
    assert chunks[0].chunk_id == "my_resume.txt:0"
    assert chunks[1].chunk_id == "my_resume.txt:1"


def test_ingest_document_empty_file_returns_no_chunks(tmp_path):
    f = tmp_path / "empty.txt"
    f.write_text("")
    assert ingest_document(f) == []


def test_ingest_directory_processes_multiple_files(tmp_path):
    (tmp_path / "resume.txt").write_text("Real content about deployed systems and testing.")
    (tmp_path / "portfolio.md").write_text("Real content about a computer vision project I shipped.")
    (tmp_path / "notes.png").write_bytes(b"unsupported binary content")

    chunks = ingest_directory(str(tmp_path))

    sources = {c.chunk_id.split(":")[0] for c in chunks}
    assert "resume.txt" in sources
    assert "portfolio.md" in sources
    assert "notes.png" not in sources  # unsupported type silently skipped, not crashed on
